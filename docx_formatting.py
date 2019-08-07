from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import copy
import re
import os
import csv
import pandas as pd

#Global variables
originDirectory = "C:/Users/rdmcl/Box Sync/Python_files/Word_Integration/original"
saveDirectory = "C:/Users/rdmcl/Box Sync/Python_files/Word_Integration/parsed"
keywords = {
    "some": WD_COLOR_INDEX.RED,
    "first": WD_COLOR_INDEX.YELLOW,
    "No": WD_COLOR_INDEX.RED,
    "know": WD_COLOR_INDEX.YELLOW
}
counts = {}
columnNames = ['Participant ID']
pIDs = []


def find_occurences_in_paragraph(paragraph, search):
    return [m.start() for m in re.finditer(search, paragraph.text)]

def apply_format_to_range(paragraph, start, end, format_func, search):
    for run in get_target_runs(paragraph, start, end, search):
        format_func(run)

def get_target_runs(paragraph, start, end, search): 
    speaker = paragraph.text[:5].lower()
    targets = []
    
    
    #Must be done in a while loop because splitting the run will modify
    #paragraph.runs
    i = 0
    past_start = False
    while(i < len(paragraph.runs)):
        run = paragraph.runs[i] #todo: FIND A WAY OF MAKING THIS CASE INSENSITIVE
        run_start = sum([len(r.text) for r in paragraph.runs[:i]])#inefficient but guaranteed correct
        run_end = run_start + len(run.text)
        
        run_contains_start = (run_start <= start <= run_end)
        run_contains_end = (run_start <= end <= run_end)
        #Split run in three, take middle part
        if(run_contains_start and run_contains_end):
            split_runs = split_run_in_three(paragraph, run, start-run_start, end-run_end)
            targets = [split_runs[1]]
            #print([r.text for r in targets])
            if(speaker == 'paren' or speaker == 'child'):
                counts[search][speaker] = counts[search][speaker] + 1 
            return targets
        #Split run, take second half
        elif(run_contains_start and not run_contains_end):
            past_start = True 
            split_runs = split_run_in_two(paragraph, run, start-run_start) 
            targets.append(split_runs[1]) 
            i += 1 #skip run that was added by splitting run 
        #Take whole run
        elif(past_start and not run_contains_end):
            targets.append(run)
        #Split run, take first half
        elif(past_start and run_contains_end):
            split_runs = split_run_in_two(paragraph, run, end-run_start)
            targets.append(split_runs[0])
            return targets
        i += 1
    return targets

def split_run_in_two(paragraph, run, split_index):
    index_in_paragraph = paragraph._p.index(run.element)

    text_before_split = run.text[0:split_index]
    text_after_split = run.text[split_index:]
    
    run.text = text_before_split
    new_run = paragraph.add_run(text_after_split)
    copy_format_manual(run, new_run)
    paragraph._p[index_in_paragraph+1:index_in_paragraph+1] = [new_run.element]
    return [run, new_run]

def split_run_in_three(paragraph, run, split_start, split_end):
    first_split = split_run_in_two(paragraph, run, split_end)
    second_split = split_run_in_two(paragraph, run, split_start)
    return second_split + [first_split[-1]]


def copy_format_manual(runA, runB):
    fontB = runB.font
    fontA = runA.font
    fontB.bold = fontA.bold
    fontB.italic = fontA.italic
    fontB.underline = fontA.underline
    fontB.strike = fontA.strike
    fontB.subscript = fontA.subscript
    fontB.superscript = fontA.superscript
    fontB.size = fontA.size
    fontB.highlight_color = fontA.highlight_color
    fontB.color.rgb = fontA.color.rgb
    #Probably others...

def parse_document(filename, search, color):
    doc = Document(filename)

    for paragraph in doc.paragraphs:
        #print(find_occurences_in_paragraph(paragraph, search))
        format_func = lambda x:x.font.__setattr__('highlight_color', color)
        for start in find_occurences_in_paragraph(paragraph, search):
            apply_format_to_range(paragraph, start, start + len(search), format_func, search)

    doc.save(filename)


if __name__ == '__main__':

    directory = os.fsencode(originDirectory)
    for file in os.listdir(directory):
        filename = os.fsdecode(file)
        #filename = "demo.docx"
        print(filename)
        pIDs.append(filename)
        savedName = saveDirectory + "/" + filename

        doc = Document("original/"+filename) 
        doc.save(savedName) # Creates a copy to work with and highlight
        for key in keywords:
            counts[key] = {"paren":0, "child":0}
            parse_document(savedName, key, keywords[key]) #actual algorithm is in get_target_runs
            print("\tMother count for \'" + key + "\' is: " + str(counts[key]["paren"]))
            print("\tChild count for \'" + key + "\' is: " + str(counts[key]["child"]))
            str_parent = 'Parent \'' + key + '\' counts'
            str_child = 'Child \'' + key + '\' counts'
            columnNames.append(str_parent)
            columnNames.append(str_parent)
    #print(pIDs)
    #print(columns)
    
    df = pd.DataFrame(columns=columnNames) #todo: this needs to be linked somehow with the filename and row in the dataframe.
    df['Participant ID'] = pIDs
#    for key in keywords:
#        colParent = 'Parent \'' + key + '\' counts'
#        colChild = 'Child \'' + key + '\' counts'
#        df
    print(df)
            
            # Parent
            # Child
            # Speaker 1
            # New Speaker
            
    
    
    