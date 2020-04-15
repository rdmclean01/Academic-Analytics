##############################################################
# Author: Ryan McLean
# Phone Number: 801-680-1086
# Email: rdmclean01@gmail.com

# Please, do not hesitate to call/text if you are having issues with this code
# I like to think of myself as a very approachable person
##############################################################

# This script uses libraries (see the import statements below)
#   If you have trouble, you might need to run the following commands in
#   a command window to install the packages
#       pip install python-docx
#       pip install pandas


from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
import copy
import re
import os
import csv
import pandas as pd
import tkinter as tk
from tkinter import filedialog


#Global variables
originDirectory = "NEW TRANSCRIPTIONS (Final Clean Up)"
#originDirectory = "original"
saveDirectory = "parsedFinal"
#originDirectory = "C:/Users/rdmclean/Box/Python_files/Word_Integration/original"
#saveDirectory = "C:/Users/rdmclean/Box/Python_files/Word_Integration/parsed"
filename_keys= 'NEW_keywords.csv' # Name of csv with all the words. This file must be in the current working directory
head_emotion = 'Emotion Words'  # Column name in csv for Emotion Words
head_emoter = 'Emoter Words'    # Column name in csv for Emoter Words
head_objects = 'Object Words'   # Column name in csv for Object Words

# Options for colors are at https://python-docx.readthedocs.io/en/latest/api/enum/WdColorIndex.html
#   This script will highlight all the words it finds in the respective color
color_emotion = WD_COLOR_INDEX.YELLOW
color_emoter = WD_COLOR_INDEX.RED
color_object = WD_COLOR_INDEX.GREEN
color_page = WD_COLOR_INDEX.VIOLET

counts = {} # Initialize an empty dictionary to keep track of the number of times a word is spoken by parent and child
columnNames = ['ID', 'PageNum', 'P_WordCount', 'C_WordCount'] # A list to keep track of column names for final data file
df_counts = pd.DataFrame(columns=columnNames)
countWords = False # This boolean keeps track so the words are only counted the first iteration
numPages = 30 # The number of different pictures the children could be shown
currSpeaker = 'None'
currPage = 'None'
currDoc = 'None'
currIndex = 'None'


##############################################################
# All of the following are the defined functions
# In order to follow program flow, start in main below
##############################################################


##############################################################
# find_occurences_in_paragraph uses regular expressions to search
# the paragraph text and find instances of the appropriate word

# Called by parse_Document
##############################################################
def find_occurences_in_paragraph(paragraph, search):
    # Find the start positions for each matched word in the run using regular expressions
    # returns a list of the start positions of each found iteration of the word
    return [m.start() for m in re.finditer(search.lower(), paragraph.text.lower())] 

##############################################################
# apply_format_to_range calls get_target_runs function for 
# each of the instances found by find_occurences_in_paragraph
# then it highlights the appropriate runs

# Called by parse_Document
##############################################################
def apply_format_to_range(paragraph, start, end, format_func, search):
    for run in get_target_runs(paragraph, start, end, search):
        #print("    "+search + " " + run.text)
        format_func(run) # Lambda function that actually does the highlighting

##############################################################
# get_target_runs goes through the given paragraph
# if that paragraph contains a key word, it looks for the run
# that contains that word and splits that word into a separate
# run so the appropriate highlighting can be done to the key word

# Called by apply_format_to_range
##############################################################
def get_target_runs(paragraph, start, end, search): 
    speaker = paragraph.text[:5].lower() # Flag for either the parent or the child
    targets = [] # List to hold the appropriate runs
    if(speaker != 'paren' and speaker != 'child'):
        print("Paragraph starts wrong: ",speaker)
        return targets
    #Must be done in a while loop because splitting the run will modify # paragraph.runs
    i = 0 # A count variable to keep track of the iterations
    past_start = False # Initialize this boolean
    while(i < len(paragraph.runs)):
        run = paragraph.runs[i] # Grab the current run
        run_start = sum([len(r.text) for r in paragraph.runs[:i]])# Counts the length of each previous run to find the start of the current run
        run_end = run_start + len(run.text) # Find the end of the run 

        run_contains_start = (run_start <= start <= run_end) # Boolean to see if 
        run_contains_end = (run_start <= end <= run_end)
        #Split run in three, take middle part
        if(run_contains_start and run_contains_end):
            split_runs = split_run_in_three(paragraph, run, start-run_start, end-run_end)
            targets = [split_runs[1]]
            #print('Run contains start and end')
            # print([r.text for r in targets])
            if(speaker == 'paren' or speaker == 'child'):
                counts[search][speaker] = counts[search][speaker] + 1 
            if(speaker == 'paren'):
                if(pd.isna(df_counts.loc[currIndex, 'P_' + search])):
                    #print("\t",currIndex, "\t",search," added to Parent counts NA")
                    df_counts.loc[currIndex, 'P_' + search] = 1
                else:
                    #print("\t",currIndex, "\t",search," added to Parent counts")
                    df_counts.loc[currIndex, 'P_' + search] += 1
                #print(currIndex,'Parent ',search, df_counts.loc[currIndex, 'P_' + search])
            if(speaker == 'child'):
                if(pd.isna(df_counts.loc[currIndex, 'C_' + search])):
                    #print("\t",currIndex, "\t",search," added to Child counts NA")
                    df_counts.loc[currIndex, 'C_' + search] = 1
                else:
                    #print("\t",currIndex, "\t",search," added to Child counts")
                    df_counts.loc[currIndex, 'C_' + search] += 1
                #print(currIndex,'Child ',search, df_counts.loc[currIndex, 'C_' + search])
            return targets
        #Split run, take second half
        elif(run_contains_start and not run_contains_end):
            past_start = True 
            split_runs = split_run_in_two(paragraph, run, start-run_start) 
            targets.append(split_runs[1]) 
            # print('Run contains start but not end ')
            # print([r.text for r in targets])
            i += 1 #skip run that was added by splitting run 
        #Take whole run
        elif(past_start and not run_contains_end):
            targets.append(run)
            # print('Past start and not contains end')
            # print([r.text for r in targets])
        #Split run, take first half
        elif(past_start and run_contains_end):
            split_runs = split_run_in_two(paragraph, run, end-run_start)
            targets.append(split_runs[0])
            #print('Past start and contains end')
            # print( [r.text for r in targets])
            if(speaker == 'paren'):
                if(pd.isna(df_counts.loc[currIndex, 'P_' + search])):
                    #print("\t",currIndex, "\t",search," added to Parent counts NA")
                    df_counts.loc[currIndex, 'P_' + search] = 1
                else:
                    #print("\t",currIndex, "\t",search," added to Parent counts")
                    df_counts.loc[currIndex, 'P_' + search] += 1
                #print(currIndex,'Parent ',search, df_counts.loc[currIndex, 'P_' + search])
            if(speaker == 'child'):
                if(pd.isna(df_counts.loc[currIndex, 'C_' + search])):
                    #print("\t",currIndex, "\t",search," added to Child counts NA")
                    df_counts.loc[currIndex, 'C_' + search] = 1
                else:
                    #print("\t",currIndex, "\t",search," added to Child counts")
                    df_counts.loc[currIndex, 'C_' + search] += 1
                #print(currIndex,'Child ',search, df_counts.loc[currIndex, 'C_' + search])
            return targets
        i += 1
    return targets
def split_para_in_two(paragraph, split_index, currSpeaker):
    #index_in_paragraph - paragraph._p.index()

    text_before_split = paragraph.text[0:split_index]
    text_after_split = paragraph.text[split_index:0]

    paragraph.text = text_before_split[:5] + 'timestamp' + text_after_split
    first_paragraph = paragraph.insert_paragraph_before(text_before_split) 
##############################################################
# split_run_in_two would split a certain run into two runs.
# E.g., Original Run = That is a snake!
#   Original Run = That is a 
#   New Run = snake! 

# Called from get_target_runs AND split_run_in_three
##############################################################
def split_run_in_two(paragraph, run, split_index):
    index_in_paragraph = paragraph._p.index(run.element)

    text_before_split = run.text[0:split_index] # Grab text before keyword
    text_after_split = run.text[split_index:]   # Grab text after keyword
    
    run.text = text_before_split # Update original run to only contain before split
    new_run = paragraph.add_run(text_after_split) # Create a new run with the text after the split
    copy_format_manual(run, new_run) # Maintain the formatting that was in the previous run
    
    #Put the new run in the right place in the paragraph, instead of the end 
    paragraph._p[index_in_paragraph+1:index_in_paragraph+1] = [new_run.element] 
    return [run, new_run] # Returns a list of [That is a, snake!]

##############################################################
# split_run_in_three uses split_run_in_two to split a run around a word
# E.g., Original Run = What does sadness feel like?
#   Original Run = What does 
#   1st New Run = sadness
#   2nd New Run = feel like? 

# Called from get_target_runs
# Returns an array with 3 elements [Run before key word, Run containing key word, Run after key word]
##############################################################
def split_run_in_three(paragraph, run, split_start, split_end):
    first_split = split_run_in_two(paragraph, run, split_end) #Splits 'What does sadness' from 'feel like?' 
    second_split = split_run_in_two(paragraph, run, split_start) # Splits 'What does ' from 'sadness'
    return second_split + [first_split[-1]] # Returns a list of [What does, sadness, feel like]

##############################################################
# copy_format_manual is used to copy the previous highlighting 
# of a run into a new run. This is used when a new
# run is created because of a keyword

# Called from split_run_in_two
# Returns nothing
##############################################################
def copy_format_manual(runA, runB):
    fontB = runB.font
    fontA = runA.font
    # fontB.highlight_color = fontA.highlight_color # For some reason, this highlighting does not work well when highlighting is null
    fontB.bold = fontA.bold
    fontB.italic = fontA.italic
    fontB.underline = fontA.underline
    fontB.strike = fontA.strike
    fontB.subscript = fontA.subscript
    fontB.superscript = fontA.superscript
    fontB.size = fontA.size
    fontB.color.rgb = fontA.color.rgb
    # Add here any other formatting things that your future needs require
    return

##############################################################
# parseDocument opens the saved file and then goes through each paragraph
# in the document. It calls countWords if appropriate, then
# finds the occurences and then does the highlighting. 

# Called from getCounts
# Returns nothing
##############################################################
def parse_document(filename, search, color):
    global currIndex
    doc = Document(filename)
    for paragraph in doc.paragraphs:
        # print(find_occurences_in_paragraph(paragraph, search)) # Prints starting location of matched words
        page_in_para = re.search('\[([pP]age )([\d\w]+)\]', paragraph.text)
        if(page_in_para):
            currPage = page_in_para[2]
            currIndex = currDoc + '-Page-' + currPage
            #print('Index changed to', currIndex)
            #df_counts.loc[currIndex, 'ID'] = currDoc
            #df_counts.loc[currIndex, 'PageNum'] = currPage
            #print(currIndex)
        #print('Index outside is', currIndex)
        #if(countWords): # Count the words for each paragraph the first time it is parsed
        #    updateWordCount(paragraph.text[:5].lower(), paragraph, df, filename) # Call the updateWordCounts function

        # Define a lamda function to do the highlighting
        format_func = lambda x:x.font.__setattr__('highlight_color', color)
        
        # This is slightly more complicated, 
        #   First, call the find_occurences_in_paragraph to find the all the starting locations of the key word in the paragraph (as a list)
        #   Second, call the apply_format_to_range to highlight those instances of the key word
        for start in find_occurences_in_paragraph(paragraph, search):
            # print(start)
            length = len(search)
            if(search[:5] == '[^\w]'): # Certain words like he, she, her require additional regex to not grab heard, shepherd, etc. 
                length = len(search) - 9   # So, just shorten this highlighting length for simplification. 
            if(search[:6] == '\[page'): # Same thing here
                length = len(search) - 3 # Shorten off the \ characters included for regex
            apply_format_to_range(paragraph, start, start + length, format_func, search)

    doc.save(filename) # Save the updated, highlighted, and counted file
    return

##############################################################
# getCounts calls parse_Document for each key word and then 
# updates the dataframe with the found words

# Called from main for each key word (emotions, emoters, objects)
# Returns nothing
##############################################################
def getCounts(savedName, key_emotion, color, pageNum):
    global currIndex
    # print(key_emotion) # Update the user with the current word to show progress
    if(pageNum):
        counts[key_emotion] = {"paren":0, "child":0} # Initialize the counts for current word at zero
        parse_document(savedName, key_emotion, color) # Call the function parse_Document
    else:
        key_long = '\[page ' + key_emotion + '\]'
        counts[key_long] = {"paren":0, "child":0} # Initialize the counts for current word at zero
        parse_document(savedName, key_long, color) # Call the function parse_Document
    return

##############################################################
# highlightAll will go through and highlight everything one
# color. Used to initalize everything to white at beginning

# Called from main to highlight everything white
# Returns nothing
##############################################################
def highlightAll(filename, color):
    doc = Document(filename) # opens the document
    for paragraph in doc.paragraphs: 
        # print(find_occurences_in_paragraph(paragraph, search))

        # Define a lamda function to do the highlighting
        format_func = lambda x:x.font.__setattr__('highlight_color', color)
        # Perform highlighting on the paragraph runs
        for run in paragraph.runs:
            format_func(run)
    return

##############################################################
# updateWordCount counts the number of words in the 
# given paragraph, and updates the dataframe's wordcount

# Called from ParseDocument
# Returns: Nothing
##############################################################
def updateWordCount(filename, currIndex):
    doc = Document(filename)
    for paragraph in doc.paragraphs:
        speaker = paragraph.text[:5].lower()
        #print('\t',paragraph.text)
        page_in_para = re.search('\[([pP]age )([\d\w]+)\]', paragraph.text)
        if(page_in_para):
            currPage = page_in_para[2]
            currIndex = currDoc + '-Page-' + currPage
            df_counts.loc[currIndex, 'ID'] = currDoc
            df_counts.loc[currIndex, 'PageNum'] = currPage
        #head, tail = os.path.split(filename)
        if(pd.isna(df_counts.loc[currIndex, 'P_WordCount'])):
            df_counts.loc[currIndex, 'P_WordCount'] = 0
            df_counts.loc[currIndex, 'C_WordCount'] = 0
        currParent = df_counts.loc[currIndex, 'P_WordCount']
        currChild = df_counts.loc[currIndex, 'C_WordCount']
        if(paragraph.text == ""):
            words = 0
        else:
            wordcount = len(re.findall("(\S+)", paragraph.text)) - 1 #Count the total number of words minus 1 for the speaker and 1 for the timestamp
            # print('    Word count: ' + str(wordcount))
            pageNum = len(re.findall(r'\[[pP]age [\d\w]+\]', paragraph.text)) # Find all the [page ##] to subtract from word count
            wordCorrect = len(re.findall(r'\[\[\w+\]\]', paragraph.text)) # Find all the instances of 'nake [[snake]] where word is corrected
            name = len(re.findall(r'\[[a-zA-Z\']+ name\]', paragraph.text))
            words = wordcount - (2*pageNum) - wordCorrect - name #Each [page ##] has two words, each [[corrected]] has one, and count [child's name] as 1 word instead of two
            # print("Page Numbers finds: " + str(pageNum))
            # print("wordCorrect finds: " + str(wordCorrect))
            # print("Names: " + str(name))
        if(speaker == 'paren'):
            df_counts.loc[currIndex, 'P_WordCount'] = currParent + words
            # print('        Added Parent ' + str(words) + ' for a total of ' + str(df_counts.loc[currIndex, 'P_WordCount']))
        elif(speaker == 'child'):
            df_counts.loc[currIndex, 'C_WordCount'] = currChild + words
            # print('        Added Child ' + str(words) + ' for a total of ' + str(df_counts.loc[currIndex, 'C_WordCount']))
    return


##############################################################
# This is where the program will actually start
# Then, it will follow through the function calls
##############################################################
if __name__ == '__main__':
    root = tk.Tk()
    root.withdraw()
    fileName_keys = filedialog.askopenfile(title = "Select a csv file containing your keywords",filetypes = ( ("CSV Files", "*.csv"), ("All Files", "*.*")) )
    originDirectory = filedialog.askdirectory(title = "Select the folder with your original files")
    saveDirectory = filedialog.askdirectory(title = " Select the folder where you want to parsed files")
    keywords = pd.read_csv(filename_keys, delimiter=',') #Read in key words from the csv file
    ##############################################################
    # Dynamically create the column names for the final data file
    # using all the given words in the current csv file

    # IMPORTANT: If at some point, you add a new column of words,
    #   You will need to come in here and duplicate this code for your column
    ##############################################################
    for emotion in keywords.loc[:,head_emotion]: # First do the emotions
        if(isinstance(emotion,str)):
            df_counts.loc[currIndex, 'P_' + emotion] = 0
            df_counts.loc[currIndex, 'C_' + emotion] = 0
    for emoter in keywords.loc[:,head_emoter]:  # Then do the emoters
        if(isinstance(emoter, str)):
            df_counts.loc[currIndex, 'P_' + emoter] = 0
            df_counts.loc[currIndex, 'C_' + emoter] = 0
    for referent in keywords.loc[:,head_objects]:
        if(isinstance(referent, str)):          # Then do the objects
            df_counts.loc[currIndex, 'P_' + referent] = 0
            df_counts.loc[currIndex, 'C_' + referent] = 0
    # Now, create a dataframe to hold your data with the appropriate column names
    #df_counts = pd.DataFrame(columns=columnNames)


    ### This directory section will make the code execute on each of the files in the originDirectory
    directory = os.fsencode(originDirectory)
    for file in os.listdir(directory): # Grab each file, one at a time
        filename = os.fsdecode(file) # Take the code from computer language to a path
        currDoc = filename[:(len(filename)-5)]
        print('Working to analyze ' + currDoc) # User-friendly message
        savedName = saveDirectory + "/" + filename # I want to leave the original files alone, so put edited files in a different location

        doc = Document(originDirectory +'/'+filename) # Open the original file
        doc.save(savedName) # Creates a copy to work with and highlight (basically a Save As)
        #updateWordCount(df_counts, savedName, currIndex)
        df_counts.loc[currIndex, 'P_WordCount'] = 0 # Initialize the Parent word count for the file
        df_counts.loc[currIndex, 'C_WordCount'] = 0 # Initialize the Child word count for the file
        updateWordCount(savedName, currIndex)

        # Go through each row of the key words you read in from the file
        for index, row in keywords.iterrows():
            # Since not all rows are the same length, the if(isinstance(...)) only executes the code on cells with words
            # getCounts is a function (look for implentation above to see where program goes from here)
            if(isinstance(row[head_emotion],str)):
                getCounts(savedName,row[head_emotion], color_emotion,True) # Search for the emotion words in the file text
            if(isinstance(row[head_emoter],str)):
                getCounts(savedName,row[head_emoter], color_emoter,True) # Search for the emoter words in the file text
            if(isinstance(row[head_objects],str)):
                getCounts(savedName,row[head_objects], color_object,True) # Search for the object words in the file text
        df_counts.to_csv(r'parsedFinal\word_counts_final.csv',index=True, index_label='Participant_ID')

    df_counts.fillna(0, inplace=True)
    # Once you are done, write your dataframe to a csv so you can run your statistics :)
    df_counts.to_csv(r'parsedFinal\word_counts_final.csv',index=True, index_label='Participant_ID')
    
    # Give the user a message to let them know they are done. 
    print("\n\n\nOutput for each file successfully written to csv file in parsed folder\n\n\n")

    # Create a smaller file
    df_final = pd.DataFrame(df_counts, columns=['ID','PageNum','P_WordCount','C_WordCount'])
    #print(df_final.head())
    # End of Program
