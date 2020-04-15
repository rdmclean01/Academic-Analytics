##############################################################
# Author: Ryan McLean
# Phone Number: 801-680-1086
# Email: rdmclean01@gmail.com

# Please, do not hesitate to call/text if you are having issues with this code
# I like to think of myself as a very approachable person
##############################################################

# This script uses libraries (see the import statements below)
#   If you have trouble, you might need to run the following commands in
#   a command window to install the packages
#       pip install python-docx
#       pip install pandas


from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
import copy
import re
import os
import csv
import pandas as pd
import tkinter as tk
from tkinter import filedialog


#Global variables
originDirectory = "small" # Default will get rewritten by dialog menu
saveDirectory = "parsedFinal"                            # Default will get rewritten by dialog menu
filename_keys= 'New keywords.csv' # Default will get rewritten by dialog menu
head_emotion = 'Emotion Words'  # Column name in csv for Emotion Words
head_emoter = 'Emoter Words'    # Column name in csv for Emoter Words
head_objects = 'Object Words'   # Column name in csv for Object Words

counts = {} # Initialize an empty dictionary to keep track of the number of times a word is spoken by parent and child
columnNames = ['ID', 'PageNum', 'P_transcript', 'C_transcript', 'P_WordCount', 'C_WordCount'] # A list to keep track of column names for final data file
df_counts = pd.DataFrame(columns=columnNames)
countWords = False # This boolean keeps track so the words are only counted the first iteration
numPages = 30 # The number of different pictures the children could be shown
currSpeaker = 'None'
currPage = 'None'
currDoc = 'None'
currIndex = 'None'



##############################################################
# updateWordCount counts the number of words in each paragraph

# Called from Main
# Returns: Nothing
##############################################################
def updateWordCount(filename):
    global currIndex    # Reference the global current Index
    global currSpeaker  # Reference the global current Speaker
    global currDoc      # FIXME
    doc = Document(filename) # Open the current document
    for paragraph in doc.paragraphs:
        # print('\t',paragraph.text)
        
        #Update from a speaker change
        if(paragraph.text[:5].lower() == 'paren' or paragraph.text[:5].lower() == 'child'):
            currSpeaker = paragraph.text[:5].lower()

        #Update from a page change
        page_in_para = re.search('\[([pP]age )([\d\w]+)\]', paragraph.text)
        if(page_in_para):
            currPage = page_in_para[2]
            currIndex = currDoc + '-Page-' + currPage
            if(currPage.lower == 'stop'):
                print('Stop Page Detected')
            else:
                df_counts.loc[currIndex, 'ID'] = currDoc
                df_counts.loc[currIndex, 'PageNum'] = currPage

        # Save the transcript
        if((currSpeaker == 'paren') & (len(re.findall(r'\d\d:\d\d', paragraph.text)) == 0)):
            if(pd.isna(df_counts.loc[currIndex, 'P_transcript'])):
                df_counts.loc[currIndex, 'P_transcript'] = ""
            df_counts.loc[currIndex, 'P_transcript'] = df_counts.loc[currIndex, 'P_transcript'] + ' ' + paragraph.text
        elif((currSpeaker == 'child') & (len(re.findall(r'\d\d:\d\d', paragraph.text)) == 0)):
            if(pd.isna(df_counts.loc[currIndex, 'C_transcript'])):
                df_counts.loc[currIndex, 'C_transcript'] = ""
            df_counts.loc[currIndex, 'C_transcript'] = df_counts.loc[currIndex, 'C_transcript'] + ' ' + paragraph.text

        #Handle the initialization of a new cell
        if(pd.isna(df_counts.loc[currIndex, 'P_WordCount'])):
            df_counts.loc[currIndex, 'P_WordCount'] = 0
            df_counts.loc[currIndex, 'C_WordCount'] = 0
        currParent = df_counts.loc[currIndex, 'P_WordCount'] # Grab the current sum of parent words in paragraph
        currChild = df_counts.loc[currIndex, 'C_WordCount']  # Grab the current sum of child  words in paragraph
        
        # Update the Word Count
        if(paragraph.text == ""):
            words = 0
        else:
            wordcount = len(re.findall("(\S+)", paragraph.text)) # Count the total number of words
            speakWords = 2*len(re.findall(r'\d\d:\d\d', paragraph.text)) # Take off two words every time a new speaker is referenced with a time stamp
            pageNum = len(re.findall(r'\[[pP]age [\d\w]+\]', paragraph.text)) # Find all the [page ##] to subtract from word count
            wordCorrect = len(re.findall(r'\[\[\w+\]\]', paragraph.text)) # Find all the instances of 'nake [[snake]] where word is corrected
            name = len(re.findall(r'\[[a-zA-Z\']+ name\]', paragraph.text)) # Find the instances where name has been de-identified
            words = wordcount - speakWords - (2*pageNum) - wordCorrect - name #Each [page ##] has two words, each [[corrected]] has one, and count [child's name] as 1 word instead of two
        if(currSpeaker == 'paren'):
            df_counts.loc[currIndex, 'P_WordCount'] = currParent + words
            # print('        Added Parent ' + currIndex + ' ',words,' for a total of ' + str(df_counts.loc[currIndex, 'P_WordCount']))
        elif(currSpeaker == 'child'):
            df_counts.loc[currIndex, 'C_WordCount'] = currChild + words
            # print('        Added Child ' + currIndex + ' ',words,' for a total of ' + str(df_counts.loc[currIndex, 'C_WordCount']))
    return





##############################################################
# This is where the program will actually start
# Then, it will follow through the function calls
##############################################################
if __name__ == '__main__':
    root = tk.Tk()
    root.withdraw()
    fileName_keys = filedialog.askopenfile(title = "Select a csv file containing your keywords",filetypes = ( ("CSV Files", "*.csv"), ("All Files", "*.*")) )
    originDirectory = filedialog.askdirectory(title = "Select the folder with your original files")
    saveDirectory = filedialog.askdirectory(title = " Select the folder where you want to the output file to be")
    keywords = pd.read_csv(filename_keys, delimiter=',') #Read in key words from the csv file
    ##############################################################
    # Dynamically create the column names for the final data file
    # using all the given words in the current csv file

    # IMPORTANT: If at some point, you add a new column of words,
    #   You will need to come in here and duplicate this code for your column
    ##############################################################
    for emotion in keywords.loc[:,head_emotion]: # First do the emotions
        if(isinstance(emotion,str)):
            df_counts.loc[currIndex, 'P_' + emotion] = 0
            df_counts.loc[currIndex, 'C_' + emotion] = 0
    for emoter in keywords.loc[:,head_emoter]:  # Then do the emoters
        if(isinstance(emoter, str)):
            df_counts.loc[currIndex, 'P_' + emoter] = 0
            df_counts.loc[currIndex, 'C_' + emoter] = 0
    for referent in keywords.loc[:,head_objects]:
        if(isinstance(referent, str)):          # Then do the objects
            df_counts.loc[currIndex, 'P_' + referent] = 0
            df_counts.loc[currIndex, 'C_' + referent] = 0


    ### This directory section will make the code execute on each of the files in the originDirectory
    directory = os.fsencode(originDirectory)
    for file in os.listdir(directory): # Grab each file, one at a time
        filename = os.fsdecode(file) # Take the code from computer language to a path
        currDoc = filename[:(len(filename)-5)] # Grab the name of the current document from the full path
        print('Working to analyze ' + currDoc) # User-friendly message
        openName = originDirectory + "/" + filename # I want to leave the original files alone, so put edited files in a different location

        doc = Document(originDirectory +'/'+filename) # Open the original file
        df_counts.loc[currIndex, 'P_WordCount'] = 0 # Initialize the Parent word count for the file
        df_counts.loc[currIndex, 'C_WordCount'] = 0 # Initialize the Child word count for the file
        updateWordCount(openName)
        df_counts.to_csv(saveDirectory + r'\word_counts_final.csv',index=True, index_label='Participant_ID')

        # indexNames = df_counts[df_counts['PageNum'] == 'stop']
        # df_counts.drop(indexNames, inplace=True)

        # Go through each row of the key words you read in from the file
    print('Analyzing Key Words')
    for index, row in keywords.iterrows():
        print(row[head_emotion], ' ', row[head_emoter],' ',row[head_objects])
        for indexCount, rowCount in df_counts.iterrows():
            p_transcript = str(df_counts.loc[indexCount, 'P_transcript']).lower()
            c_transcript = str(df_counts.loc[indexCount, 'C_transcript']).lower()
            print (p_transcript)
            print (c_transcript)
            if(isinstance(row[head_emotion],str)):
                parentWord = len(re.findall(row[head_emotion], p_transcript))
                childWord =  len(re.findall(row[head_emotion], c_transcript))
                df_counts.loc[indexCount, 'P_' + row[head_emotion]] = parentWord
                df_counts.loc[indexCount, 'C_' + row[head_emotion]] = childWord
            if(isinstance(row[head_emoter],str)):
                parentWord = len(re.findall(row[head_emoter], p_transcript))
                childWord =  len(re.findall(row[head_emoter], c_transcript))
                df_counts.loc[indexCount, 'P_' + row[head_emoter]] = parentWord
                df_counts.loc[indexCount, 'C_' + row[head_emoter]] = childWord
            if(isinstance(row[head_objects],str)):
                parentWord = len(re.findall(row[head_objects], p_transcript))
                childWord =  len(re.findall(row[head_objects], c_transcript))
                df_counts.loc[indexCount, 'P_' + row[head_objects]] = parentWord
                df_counts.loc[indexCount, 'C_' + row[head_objects]] = childWord

    df_counts.fillna(0, inplace=True)
    # Once you are done, write your dataframe to a csv so you can run your statistics :)
    df_counts.to_csv(saveDirectory + r'\word_counts_final.csv',index=True, index_label='Participant_ID')
    
    # Give the user a message to let them know they are done. 
    print("\n\n\nOutput for each file successfully written to csv file in parsed folder\n\n\n")

    # End of Program
